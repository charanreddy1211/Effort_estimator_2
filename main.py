# import os
# import random
# import tempfile
# from fastapi import FastAPI, File, UploadFile, HTTPException
# from fastapi.responses import JSONResponse
# from fastapi.middleware.cors import CORSMiddleware
# from typing import List, Dict, Any
# from dotenv import load_dotenv
# import openai

# # PDF and DOCX handling
# try:
#     from llama_index.core import SimpleDirectoryReader
# except ImportError:
#     SimpleDirectoryReader = None
# try:
#     import docx
# except ImportError:
#     docx = None

# app = FastAPI()

# # Allow CORS for frontend
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # Load environment variables
# load_dotenv()
# AZURE_OPENAI_ENDPOINT = os.getenv("Endpoint_URL") or os.getenv("ENDPOINT_URL") or os.getenv("AZURE_OPENAI_ENDPOINT")
# AZURE_OPENAI_KEY = os.getenv("API_Key") or os.getenv("API_KEY") or os.getenv("AZURE_OPENAI_KEY")
# AZURE_OPENAI_DEPLOYMENT = os.getenv("Deployment_Name") or os.getenv("DEPLOYMENT_NAME") or os.getenv("AZURE_OPENAI_DEPLOYMENT")
# AZURE_OPENAI_MODEL = os.getenv("Model_Name") or os.getenv("MODEL_NAME") or os.getenv("AZURE_OPENAI_MODEL")
# AZURE_OPENAI_API_VERSION = os.getenv("API_Version") or os.getenv("API_VERSION") or os.getenv("AZURE_OPENAI_API_VERSION")

# # Set up OpenAI client
# from openai import AzureOpenAI
# client = AzureOpenAI(
#     api_key=AZURE_OPENAI_KEY,
#     api_version=AZURE_OPENAI_API_VERSION,
#     azure_endpoint=AZURE_OPENAI_ENDPOINT.split("/openai")[0] if AZURE_OPENAI_ENDPOINT else None
# )

# # Helper: extract text from file
# async def extract_text(file: UploadFile) -> str:
#     ext = file.filename.lower().split(".")[-1]
#     if ext == "txt":
#         content = await file.read()
#         return content.decode("utf-8", errors="ignore")
#     elif ext == "pdf":
#         if not SimpleDirectoryReader:
#             raise HTTPException(500, "llama-index not installed")
#         with tempfile.TemporaryDirectory() as tmpdir:
#             path = os.path.join(tmpdir, file.filename)
#             with open(path, "wb") as f:
#                 f.write(await file.read())
#             docs = SimpleDirectoryReader(tmpdir).load_data()
#             return "\n".join([d.text for d in docs])
#     elif ext == "docx":
#         if not docx:
#             raise HTTPException(500, "python-docx not installed")
#         with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
#             tmp.write(await file.read())
#             tmp.flush()
#             doc = docx.Document(tmp.name)
#             text = "\n".join([p.text for p in doc.paragraphs])
#             os.unlink(tmp.name)
#             return text
#     else:
#         raise HTTPException(400, "Unsupported file type")

# # Helper: call GPT-4o for summarization or scenario generation
# def call_gpt(messages: List[Dict[str, str]]) -> str:
#     response = client.chat.completions.create(
#         model=AZURE_OPENAI_DEPLOYMENT,
#         messages=messages
#     )
#     return response.choices[0].message.content

# # Helper: generate architect tasks
# async def generate_tasks(scenario: str) -> List[Dict[str, Any]]:
#     prompt = f"""
# You are a Technical Lead.

# For the following project scenario, break down the work into **high-level development tasks**, such as UI Development, UX Design, API Implementation, Unit Testing, Deployment, etc. 
# and group similar tasks together. Each task should be concise, only two words, and should not repeat the instruction. tasks should be grouped like UI Development, UX Design, API Implementation, Unit Testing, Deployment, etc.
# All UI tasks or related should be grouped together, all UX tasks should be grouped together, etc.
# All UX tasks or related should be grouped together, all API tasks should be grouped together, etc.
# All ApI tasks related should be grouped together, all Testing tasks like Unit Testing, Integration Testing, System Testing ..etc should be grouped together, etc. 
# Same way for Deployment tasks or related tasks, all DevOps tasks should be grouped together, etc.

# For each task, return a **Python dictionary** with the following keys:
# - `task`: The name of the overall task (only two words, do not repeat this instruction) and group similar tasks together give single name for it
# - `hours`: Estimated realistic working hours based on task complexity (do not generate randomly), important feature to be realistic and not too high
# - `justification`: A concise explanation of the estimation and why the task is required
# - `confidence`: A realistic confidence score between 85% and 95%

# ⛔ Do NOT include summaries, titles, or extra text.  
# ✅ Only return a **list of such dictionaries in valid JSON format**.

# 📌 Example format:

# ```json
# [
#   {{
#     "task": "UI Development",
#     "hours": 16,
#     "justification": "Four screens, average four hours each",
#     "confidence": "92%"
#   }},
#   {{
#     "task": "UX Design",
#     "hours": 8,
#     "justification": "Medium-complexity wireframes",
#     "confidence": "90%"
#   }}
# ]


# Scenario:
# {scenario}
# Important: Give two words of the task only, do not repeat the text.
# Stricly: Don't give too much hours, give realistic hours based on the task complexity.
# Important and Stricly: remove the markdown code block from the response, just return the JSON list.
# """
#     messages = [
#         {"role": "system", "content": "You are a solution architect."},
#         {"role": "user", "content": prompt}
#     ]
#     response = call_gpt(messages)

#     cleaned_response = response.strip()
#     if cleaned_response.startswith("```"):
#         cleaned_response = cleaned_response.split("\n", 1)[1]
#     if cleaned_response.endswith("```"):
#         cleaned_response = cleaned_response.rsplit("\n", 1)[0]
#     cleaned_response = cleaned_response.strip()
#     # Try to parse as JSON, else return as string
#     import json
#     try:
#         return json.loads(response)
#     except Exception:
#         return response

# @app.post("/upload")
# async def upload(file: UploadFile = File(...)):
#     try:
#         text = await extract_text(file)
#     except Exception as e:
#         return JSONResponse(status_code=400, content={"error": str(e)})

#     # Step 1: Summarize and generate scenario
#     scenario_prompt = f"""
# You are a senior full-stack developer. Given the following extracted text from a document, reframe it as professional-level development instructions, suitable for a full-stack developer. Be detailed and extend the scenario as needed.

# Extracted text:
# {text}
# """
    
    
#     messages = [
#         {"role": "system", "content": "You are a helpful assistant."},
#         {"role": "user", "content": scenario_prompt}
#     ]
#     scenario = call_gpt(messages)

#     # Step 2: Generate architect tasks
#     tasks = await generate_tasks(scenario)
#     return JSONResponse(content=tasks)

# if __name__ == "__main__":
#     import uvicorn
#     uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)


import os
import json
import tempfile
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Dict, Any, Optional
from dotenv import load_dotenv
import openai

# PDF and DOCX handling
try:
    from llama_index.core import SimpleDirectoryReader
except ImportError:
    SimpleDirectoryReader = None
try:
    import docx
except ImportError:
    docx = None

app = FastAPI()

# Allow CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load environment variables
load_dotenv()
AZURE_OPENAI_ENDPOINT = os.getenv("Endpoint_URL") or os.getenv("ENDPOINT_URL") or os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_KEY = os.getenv("API_Key") or os.getenv("API_KEY") or os.getenv("AZURE_OPENAI_KEY")
AZURE_OPENAI_DEPLOYMENT = os.getenv("Deployment_Name") or os.getenv("DEPLOYMENT_NAME") or os.getenv("AZURE_OPENAI_DEPLOYMENT")
AZURE_OPENAI_MODEL = os.getenv("Model_Name") or os.getenv("MODEL_NAME") or os.getenv("AZURE_OPENAI_MODEL")
AZURE_OPENAI_API_VERSION = os.getenv("API_Version") or os.getenv("API_VERSION") or os.getenv("AZURE_OPENAI_API_VERSION")

# Set up OpenAI client
from openai import AzureOpenAI
client = AzureOpenAI(
    api_key=AZURE_OPENAI_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
    azure_endpoint=AZURE_OPENAI_ENDPOINT.split("/openai")[0] if AZURE_OPENAI_ENDPOINT else None
)

# Helper: extract text from file
async def extract_text(file: UploadFile) -> str:
    ext = file.filename.lower().split(".")[-1]
    if ext == "txt":
        content = await file.read()
        return content.decode("utf-8", errors="ignore")
    elif ext == "pdf":
        if not SimpleDirectoryReader:
            raise HTTPException(500, "llama-index not installed")
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, file.filename)
            with open(path, "wb") as f:
                f.write(await file.read())
            docs = SimpleDirectoryReader(tmpdir).load_data()
            return "\n".join([d.text for d in docs])
    elif ext == "docx":
        if not docx:
            raise HTTPException(500, "python-docx not installed")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(await file.read())
            tmp.flush()
            doc = docx.Document(tmp.name)
            text = "\n".join([p.text for p in doc.paragraphs])
            os.unlink(tmp.name)
            return text
    else:
        raise HTTPException(400, "Unsupported file type")

# Helper: call GPT-4o for summarization or scenario generation
def call_gpt(messages: List[Dict[str, str]]) -> str:
    response = client.chat.completions.create(
        model=AZURE_OPENAI_DEPLOYMENT,
        messages=messages
    )
    return response.choices[0].message.content

# Helper: generate tasks based on all inputs
async def generate_tasks(scenario: str, cloud_provider: str, persona: str) -> List[Dict[str, Any]]:
    persona_instructions = {
        "developer": "You are a Senior Developer. Provide detailed technical implementation tasks.",
        "tech_lead": "You are a Technical Lead. Provide architectural tasks and oversight.",
        "pm": "You are a Project Manager. Provide high-level project tasks and timelines."
    }
    
    provider_instructions = {
        "Azure": "The solution will be deployed on Microsoft Azure.",
        "AWS": "The solution will be deployed on Amazon Web Services.",
        "GCP": "The solution will be deployed on Google Cloud Platform."
    }
    
    prompt = f"""
{persona_instructions.get(persona, "You are a Technical Lead.")}
{provider_instructions.get(cloud_provider, "The solution will be deployed on a cloud platform.")}

For the following project scenario, break down the work into high-level development tasks:

1. Group similar tasks together (UI, UX, API, Testing, Deployment, etc.)
2. For each task, return a Python dictionary with:
   - `task`: Concise name (2-3 words),and group similar tasks together. Each task should be concise, only two words, and should not repeat the instruction. tasks should be grouped like UI Development, UX Design, API Implementation, Unit Testing, Deployment, etc.
   - `hours`: Realistic hours based on {persona} perspective
   - `justification`: Brief explanation
   - `confidence`: Confidence score (85-95%)

Important:
- Adjust estimates based on {persona} role
- Consider {cloud_provider} specifics for deployment tasks
- Be realistic with hours
- Return only JSON list, no markdown
- and group similar tasks together. Each task should be concise, only two words, and should not repeat the instruction. tasks should be grouped like UI Development, UX Design, API Implementation, Unit Testing, Deployment, etc.
- All UI tasks or related should be grouped together, all UX tasks should be grouped together, etc.
- All UX tasks or related should be grouped together, all API tasks should be grouped together, etc.
- All ApI tasks related should be grouped together, all Testing tasks like Unit Testing, Integration Testing, System Testing ..etc should be grouped together, etc. 
- Same way for Deployment tasks or related tasks, all DevOps tasks should be grouped together, etc.

Example format:
```json
[
  {{
    "task": "UI Development",
    "hours": 16,
    "justification": "Four screens at 4 hours each",
    "confidence": 0.9
  }}
]

Scenario:
{scenario}
"""

    messages = [
        {"role": "system", "content": "You are a solution architect."},
        {"role": "user", "content": prompt}
    ]
    response = call_gpt(messages)

    # Clean and parse response
    cleaned_response = response.strip()
    if cleaned_response.startswith("```json"):
        cleaned_response = cleaned_response[7:-3].strip()
    elif cleaned_response.startswith("```"):
        cleaned_response = cleaned_response.split("\n", 1)[1].rsplit("\n", 1)[0]
    
    try:
        return json.loads(cleaned_response)
    except json.JSONDecodeError:
        try:
            # Try to find JSON within the response
            start = cleaned_response.find('[')
            end = cleaned_response.rfind(']') + 1
            return json.loads(cleaned_response[start:end])
        except Exception as e:
            raise HTTPException(500, f"Failed to parse AI response: {str(e)}")

@app.post("/estimate")
async def estimate(
    file: Optional[UploadFile] = File(None),
    text: Optional[str] = Form(None),
    cloud_provider: str = Form(...),
    persona: str = Form(...)
):
    # Validate inputs
    if not file and not text:
        raise HTTPException(400, "Either file or text input is required")
    
    if persona not in ["developer", "tech_lead", "pm"]:
        raise HTTPException(400, "Invalid persona specified")
    
    if cloud_provider not in ["Azure", "AWS", "GCP"]:
        raise HTTPException(400, "Invalid cloud provider specified")

    try:
        # Get input text
        input_text = ""
        if file:
            input_text = await extract_text(file)
        elif text:
            input_text = text

        # Step 1: Summarize and generate scenario
        scenario_prompt = f"""
Reframe this input as professional development instructions for a {persona}.
Consider deployment on {cloud_provider}.

Input:
{input_text}
"""
        
        messages = [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": scenario_prompt}
        ]
        scenario = call_gpt(messages)

        # Step 2: Generate tasks
        tasks = await generate_tasks(scenario, cloud_provider, persona)
        return JSONResponse(content={"estimates": tasks})
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Processing failed: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
